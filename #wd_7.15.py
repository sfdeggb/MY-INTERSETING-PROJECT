#wd_7.15

### 使用python-docx库俩操作word文档

from docx import Document

#创建word 文档
mydoc= Document()
#print(dir(mydoc))
## 分节符（为了设置格式用的）和分页符（开始新的一页）是不同的
#标题本身及就是段落


### 设置页眉和页脚
#操作xml文档

from docx import Document
from docx import oxml#这个是来操作xml文档的

def create_element(name):
    return oxml.OxmlElement(name)
def create_attribute(element, name, value):
    element.set(oxml.ns.qn(name), value)
def add_page_number(run):
    fldChar1=create_element('w:fldchar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText=create_element("w:instrText")
    create_attribute(instrText, 'xml.space', 'preserve')
    instrText.text='PAGE'
    fldChar2=create_element()
    create_attribute(element, name, value)
    run._r.append()
    run._r.append()
    run._r.append()
    
doc = Document()
add_page_number(doc.section[0],add_run())
for i in range(10):
    doc.add_page_break()

doc.save('lujnin')